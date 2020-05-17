# !python3

from docx.enum.text import WD_COLOR_INDEX
import docx

def hightlight(phrase):
    for paragraph in document.paragraphs:
        if phrase in paragraph.text:
            for run in paragraph.runs:
                if phrase in run.text:
                    x = run.text.split(phrase)
                    run.clear()
                    for i in range(len(x)-1):
                        run.add_text(x[i])
                        run.add_text(phrase)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

"""Simple Usage of Bayoo DOCX
document = docx.Document()

paragraph1 = document.add_paragraph('text') # create new paragraph

comment = paragraph.add_comment('comment',author='Obay Daba',initials= 'od') # add a comment on the entire paragraph

paragraph2 = document.add_paragraph('text') # create another paragraph

run = paragraph2.add_run('texty') add a run to the paragraph

run.add_comment('comment') # add a comment only for the run text

paragraph.add_footnote('footnote text') # add a footnote"""