"""Usage : python /home/you/script/this.py file.docx 'phrase to highlight'  # notice the quotes
"""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import sys

source = sys.argv[1]
phrase = " ".join(sys.argv[2:]).strip("'")

doc = Document( source )

for para in doc.paragraphs :
 start = para.text.find( phrase )
 if start > -1 :
  pre = para.text[:start]
  post = para.text[start+len(phrase):]
  para.text = pre
  para.add_run(phrase)
  para.runs[1].font.highlight_color = WD_COLOR_INDEX.YELLOW
  para.add_run(post)

doc.save( source )